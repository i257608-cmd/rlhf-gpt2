"""
IEEE Conference Report Generator
===================================
Generates: results/day10/RLHF_IEEE_Report.docx
IEEE two-column conference paper format.
Authors: Muhammad Haseeb, Muavia Shakeel — FAST University Islamabad
"""

import os, json
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = "results/day10"
os.makedirs(OUT_DIR, exist_ok=True)

# ── Load results ──────────────────────────────────────────────
c1   = json.load(open("results/c1/eval_results.json"))
c2   = json.load(open("results/c2/comparison.json"))
d8   = json.load(open("results/day8/alignment_tax.json"))
sweep = pd.read_csv("results/stability_sweep/sweep_results.csv")
sweep = sweep[sweep["status"] == "ok"]

agg    = c1["aggregate"]
sft_c2 = c2["SFT (baseline)"]
ppo_c2 = c2["PPO (aligned)"]
dpo_c2 = c2["DPO (aligned)"]
opt    = d8["optimal_stopping"]
method = d8["method_alignment_tax"]

# Sweep-derived C1 averages (3 seeds, β=0.20) — used for per-seed table & body text
ppo_b02 = sweep[sweep["beta"] == 0.2]
c1_before_mean = ppo_b02["mean_reward_before"].mean()
c1_after_mean  = ppo_b02["mean_reward_after"].mean()
c1_delta_mean  = c1_after_mean - c1_before_mean
c1_improv_mean = ppo_b02["reward_improvement_pct"].mean()

sweep_grp = (
    sweep.groupby("beta")
    .agg(mean_reward=("mean_reward_after","mean"),
         std_reward =("mean_reward_after","std"),
         mean_kl    =("mean_kl_final_20","mean"),
         max_kl     =("max_kl","max"),
         collapse   =("collapse_detected","sum"))
    .reset_index()
)

# ── Document setup ────────────────────────────────────────────
doc = Document()

# Page margins (IEEE: narrow)
for section in doc.sections:
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

# ── Styles ────────────────────────────────────────────────────
styles = doc.styles

def set_style(style_name, font_name, font_size, bold=False, color=None):
    try:
        s = styles[style_name]
    except KeyError:
        s = styles.add_style(style_name, 1)
    s.font.name      = font_name
    s.font.size      = Pt(font_size)
    s.font.bold      = bold
    if color:
        s.font.color.rgb = RGBColor(*color)
    return s

# ── Helper functions ──────────────────────────────────────────
def add_heading(doc, text, level=1):
    sizes = {1: 12, 2: 10, 3: 10}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper() if level == 1 else text)
    run.bold      = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(sizes.get(level, 10))
    if level == 1:
        run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.2)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.bold      = True
    p.paragraph_format.space_after = Pt(6)

def add_table_ieee(doc, headers, rows, caption):
    """Add a formatted IEEE-style table with explicit column widths."""
    add_caption(doc, caption)
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.autofit = False

    # Force table to span full content width (8.5" - 1.5" margins = 7.0")
    content_w_twips = int(7.0 * 1440)  # 1 inch = 1440 twips
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(content_w_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Distribute columns evenly (first col slightly wider for labels)
    first_w  = Inches(7.0 / n_cols * 1.2) if n_cols > 1 else Inches(7.0)
    other_w  = Inches((7.0 - 7.0 / n_cols * 1.2) / (n_cols - 1)) if n_cols > 1 else first_w
    for ci, col in enumerate(table.columns):
        col.width = first_w if ci == 0 else other_w

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold      = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # gray background
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# TITLE BLOCK
# ═══════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(
    "Reinforcement Learning from Human Feedback on GPT-2: "
    "A Comparative Study of PPO, DPO, and KL-Penalty Stability"
)
run.bold      = True
run.font.name = "Times New Roman"
run.font.size = Pt(14)

doc.add_paragraph()

author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author_p.add_run("Muhammad Haseeb, Muavia Shakeel")
run.bold      = True
run.font.name = "Times New Roman"
run.font.size = Pt(11)

affil_p = doc.add_paragraph()
affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = affil_p.add_run(
    "Department of Computer Science\n"
    "FAST National University of Computer and Emerging Sciences\n"
    "Islamabad, Pakistan\n"
    "{haseeb, muavia}@nu.edu.pk"
)
run.font.name = "Times New Roman"
run.font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════
abs_p = doc.add_paragraph()
abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abs_run_bold = abs_p.add_run("Abstract — ")
abs_run_bold.bold      = True
abs_run_bold.font.name = "Times New Roman"
abs_run_bold.font.size = Pt(10)
abs_run = abs_p.add_run(
    f"This paper presents an empirical study of Reinforcement Learning from Human Feedback (RLHF) "
    f"applied to GPT-2 (124M parameters) on the IMDB sentiment dataset. We make three contributions: "
    f"(C1) a complete PPO-RLHF pipeline achieving a mean reward improvement from "
    f"{agg['sft_mean_reward']:.3f} to {agg['ppo_mean_reward']:.3f} "
    f"(+{agg['reward_delta']:.3f}) at an alignment tax of +{agg['alignment_tax_pct']:.1f}% perplexity; "
    f"(C2) a head-to-head comparison of PPO vs. Direct Preference Optimisation (DPO), "
    f"showing PPO is 2.5× more reward-efficient than DPO per unit perplexity increase "
    f"({method['PPO']['efficiency_reward_per_ppl_pct']:.4f} vs. "
    f"{method['DPO']['efficiency_reward_per_ppl_pct']:.4f}); and "
    f"(C3) a systematic stability sweep over KL-penalty coefficient β ∈ {{0.05, 0.10, 0.20, 0.50}} "
    f"across 3 random seeds (12 runs total), empirically identifying reward hacking onset "
    f"at β ≤ 0.10 (mean KL ≥ {opt['collapse_onset_kl']:.1f}), consistent with Gao et al. (2023). "
    f"The optimal stable configuration is β=0.20, yielding mean reward "
    f"{opt['optimal_mean_reward']:.3f} with a {opt['safety_margin_kl']:.1f}-unit "
    f"KL safety margin before collapse."
)
abs_run.font.name = "Times New Roman"
abs_run.font.size = Pt(10)
abs_p.paragraph_format.space_after = Pt(4)

kw_p = doc.add_paragraph()
kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw_bold = kw_p.add_run("Keywords — ")
kw_bold.bold      = True
kw_bold.font.name = "Times New Roman"
kw_bold.font.size = Pt(10)
kw_run = kw_p.add_run(
    "Reinforcement Learning from Human Feedback, RLHF, GPT-2, "
    "Proximal Policy Optimisation, Direct Preference Optimisation, "
    "KL Divergence, Reward Hacking, Alignment Tax, Language Models."
)
kw_run.font.name = "Times New Roman"
kw_run.font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# I. INTRODUCTION
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "I. Introduction", 1)
add_body(doc,
    "Large language models (LLMs) trained via supervised learning frequently produce outputs "
    "that are fluent but misaligned with human preferences, generating harmful, biased, or "
    "off-topic content. Reinforcement Learning from Human Feedback (RLHF) addresses this by "
    "training a reward model from human preference labels and using reinforcement learning to "
    "optimise the language model against that signal [1], [2].", indent=True)
add_body(doc,
    "A persistent challenge in RLHF is the alignment tax: reward-optimised models tend to "
    "increase perplexity as they optimise for the proxy reward. When the KL divergence "
    "penalty is too weak, the policy diverges catastrophically — a phenomenon Gao et al. "
    "(2023) term reward hacking [3]. A parallel method, Direct Preference Optimisation (DPO) "
    "[4], avoids RL entirely by reframing preference learning as a supervised classification "
    "problem, but its alignment-tax profile relative to PPO at small scale remains "
    "underexplored.", indent=True)
add_body(doc,
    "This paper makes three contributions: (C1) a full PPO-RLHF pipeline on GPT-2/IMDB with "
    "quantified alignment tax; (C2) a head-to-head PPO vs. DPO comparison on identical "
    "base model, reward model, and evaluation protocol; and (C3) a systematic sweep of the "
    "KL-penalty coefficient β identifying the reward-hacking onset threshold and recommending "
    "a stable training configuration.", indent=True)

# ═══════════════════════════════════════════════════════════════
# II. RELATED WORK
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "II. Related Work", 1)
add_body(doc,
    "Christiano et al. [1] introduced RLHF for deep RL tasks; Ouyang et al. [2] scaled it "
    "to LLMs with InstructGPT, demonstrating that RLHF substantially improves instruction "
    "following at the cost of some fluency. Bai et al. [5] applied RLHF for harmlessness and "
    "helpfulness alignment in Anthropic's Claude models.", indent=True)
add_body(doc,
    "Gao et al. [3] derived scaling laws for reward model overoptimisation, showing that the "
    "proxy reward rises then falls as KL divergence grows beyond a threshold, coining the "
    "term reward hacking. Their work motivates our C3 sweep. Rafailov et al. [4] proposed "
    "DPO as a stable, RL-free alternative; their closed-form solution avoids reward model "
    "inference at training time but introduces its own distribution-shift dynamics.", indent=True)

# ═══════════════════════════════════════════════════════════════
# III. METHODOLOGY
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "III. Methodology", 1)

add_heading(doc, "A. Base Model and Dataset", 2)
add_body(doc,
    "All experiments use GPT-2 [6] (Radford et al., 2019; 124M parameters) as the base "
    "language model and the IMDB large movie review dataset [7] (50,000 reviews) as the "
    "training and evaluation corpus. The HuggingFace Transformers 4.57.6 and TRL 0.24.0 "
    "libraries are used throughout [8]. All experiments run on CPU (Apple Silicon, "
    "MPS disabled), PyTorch 2.8.0.", indent=True)

add_heading(doc, "B. Stage 1 — Supervised Fine-Tuning (SFT)", 2)
add_body(doc,
    "GPT-2 is fine-tuned on IMDB reviews using causal language modelling (next-token "
    "prediction) to produce the SFT policy π_SFT, which serves as the initialisation "
    "point for both PPO and DPO and as the reference model for KL regularisation. "
    "Maximum sequence length: 128 tokens.", indent=True)

add_heading(doc, "C. Stage 2 — Reward Model", 2)
add_body(doc,
    "A binary sentiment reward model is trained by attaching a linear classification head "
    "to GPT-2 and fine-tuning for 3 epochs with cross-entropy loss (eval_accuracy=85.5%, "
    "eval_loss=0.356). The reward model is frozen for all subsequent stages.", indent=True)

add_heading(doc, "D. Stage 3a — PPO Training", 2)
add_body(doc,
    "PPO [9] optimises the language model policy subject to a KL divergence penalty against "
    "π_SFT. The objective is: L_PPO = E[r(x,y)] − β · KL(π_θ ‖ π_SFT). "
    "For C1, β=0.20 across 3 seeds (42, 123, 456), 200 steps. "
    "For C3, β ∈ {0.05, 0.10, 0.20, 0.50} × 3 seeds = 12 runs.", indent=True)

add_heading(doc, "E. Stage 3b — DPO Training", 2)
add_body(doc,
    "DPO [4] fine-tunes the SFT model from preference pairs without an RL loop. "
    "Preference pairs are constructed from IMDB labels (positive = preferred, "
    "negative = rejected). DPO β=0.10, 3 seeds, initialised from the SFT checkpoint.", indent=True)

add_heading(doc, "F. Evaluation Protocol", 2)
add_body(doc,
    "Models are evaluated on 10 fixed IMDB-style prompts. We report: (i) mean reward "
    "(reward model score on generated continuations); (ii) perplexity on the IMDB test "
    "set as fluency proxy; (iii) alignment tax = (PPL_aligned − PPL_SFT)/PPL_SFT × 100%; "
    "(iv) KL divergence across training; and (v) collapse detection (max KL > 20).", indent=True)

# ═══════════════════════════════════════════════════════════════
# IV. RESULTS
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "IV. Results", 1)

add_heading(doc, "A. Reward Model Performance", 2)
add_body(doc,
    "The reward model achieves 85.5% evaluation accuracy and 0.356 loss on the held-out "
    "IMDB test set after 3 epochs, confirming its reliability as a proxy reward signal.", indent=True)

add_heading(doc, "B. C1 — PPO Training Results", 2)
add_body(doc,
    f"PPO training (β=0.20, 200 steps) consistently improves mean reward across all three "
    f"seeds with no reward collapse (max KL < 15 for all runs). "
    f"Training-time reward (measured by the stability sweep protocol) increases from "
    f"{c1_before_mean:.3f} to {c1_after_mean:.3f} "
    f"(Δ={c1_delta_mean:.3f}, +{c1_improv_mean:.1f}%), consistent across all seeds (Table I). "
    f"Post-training evaluation on 10 fixed IMDB prompts (evaluate.py, seed=42) confirms "
    f"the improvement: SFT mean {agg['sft_mean_reward']:.3f} → PPO {agg['ppo_mean_reward']:.3f} "
    f"(+{agg['reward_delta']:.3f}) at a perplexity cost of +{agg['alignment_tax_pct']:.1f}%.", indent=True)

ppo_seeds = sweep[sweep["beta"] == 0.2][["seed","mean_reward_before","mean_reward_after","reward_improvement_pct","max_kl","collapse_detected"]]
add_table_ieee(doc,
    ["Seed", "Reward Before", "Reward After", "Improvement", "Max KL", "Collapse"],
    [[int(r.seed), f"{r.mean_reward_before:.3f}", f"{r.mean_reward_after:.3f}",
      f"+{r.reward_improvement_pct:.1f}%", f"{r.max_kl:.2f}", "No" if not r.collapse_detected else "Yes"]
     for r in ppo_seeds.itertuples()],
    "TABLE I.  PPO TRAINING RESULTS PER SEED (β=0.20, 200 STEPS) — Training-time reward from stability_sweep.py"
)

add_heading(doc, "C. C2 — PPO vs. DPO Comparison", 2)
add_body(doc,
    f"Both methods substantially improve reward over SFT. PPO achieves mean reward "
    f"{ppo_c2['mean_reward']:.3f} (std {ppo_c2['std_reward']:.3f}, ppl {ppo_c2['perplexity']:.1f}) "
    f"while DPO achieves {dpo_c2['mean_reward']:.3f} (std {dpo_c2['std_reward']:.3f}, "
    f"ppl {dpo_c2['perplexity']:.1f}). PPO's alignment tax (+{method['PPO']['tax_pct']:.1f}%) "
    f"is 2.4× lower than DPO's (+{method['DPO']['tax_pct']:.1f}%), giving PPO "
    f"2.5× higher reward efficiency.", indent=True)

add_table_ieee(doc,
    ["Method", "Mean Reward", "Std", "Perplexity", "Tax", "Efficiency"],
    [["SFT", f"{sft_c2['mean_reward']:.3f}", f"{sft_c2['std_reward']:.3f}", f"{sft_c2['perplexity']:.2f}", "—", "—"],
     ["PPO (β=0.20)", f"{ppo_c2['mean_reward']:.3f}", f"{ppo_c2['std_reward']:.3f}", f"{ppo_c2['perplexity']:.2f}",
      f"+{method['PPO']['tax_pct']:.1f}%", f"{method['PPO']['efficiency_reward_per_ppl_pct']:.4f}"],
     ["DPO (β=0.10)", f"{dpo_c2['mean_reward']:.3f}", f"{dpo_c2['std_reward']:.3f}", f"{dpo_c2['perplexity']:.2f}",
      f"+{method['DPO']['tax_pct']:.1f}%", f"{method['DPO']['efficiency_reward_per_ppl_pct']:.4f}"]],
    "TABLE II.  PPO VS. DPO VS. SFT (MEAN OVER 3 SEEDS)"
)

add_heading(doc, "D. C3 — KL-Penalty Stability Sweep", 2)
add_body(doc,
    f"The sweep reveals a sharp phase transition. β ≤ 0.10 produces reward hacking in all "
    f"6 runs (100% collapse rate, mean KL 17–40). β ≥ 0.20 is fully stable (0% collapse). "
    f"The hacking onset occurs at mean KL ≈ {opt['collapse_onset_kl']:.1f}, consistent "
    f"with Gao et al. [3]. The optimal configuration is β=0.20 (mean reward "
    f"{opt['optimal_mean_reward']:.3f}, safety margin {opt['safety_margin_kl']:.1f} KL units).", indent=True)

add_table_ieee(doc,
    ["β", "Mean Reward", "Std", "Mean KL", "Collapse", "Verdict"],
    [[f"{r.beta}", f"{r.mean_reward:.3f}", f"±{r.std_reward:.3f}",
      f"{r.mean_kl:.2f}", f"{int(r.collapse)}/3",
      "HACKING" if r.collapse > 0 else "stable"]
     for r in sweep_grp.itertuples()],
    "TABLE III.  STABILITY SWEEP — MEAN OVER 3 SEEDS PER β"
)

# ═══════════════════════════════════════════════════════════════
# V. DISCUSSION
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "V. Discussion", 1)

add_heading(doc, "A. Alignment Tax Trade-off", 2)
add_body(doc,
    f"The +{agg['alignment_tax_pct']:.1f}% perplexity increase for PPO confirms the "
    f"well-known reward-fluency trade-off in RLHF [2]. DPO's larger tax (+{method['DPO']['tax_pct']:.1f}%) "
    f"suggests that at small scale with automatically constructed preference pairs, "
    f"DPO distorts the output distribution more aggressively than the RL loop. "
    f"PPO's explicit KL constraint provides tighter distribution control.", indent=True)

add_heading(doc, "B. Reward Hacking and KL Threshold", 2)
add_body(doc,
    f"Our empirical threshold of KL ≈ {opt['collapse_onset_kl']:.1f} for reward hacking onset "
    f"is consistent with Gao et al.'s [3] theoretical prediction that proxy reward rises as "
    f"√KL before collapsing. The {opt['safety_margin_kl']:.1f}-unit safety margin at β=0.20 "
    f"provides a practical stopping criterion: terminate training if mean KL exceeds "
    f"{opt['collapse_onset_kl']:.1f}.", indent=True)

add_heading(doc, "C. Limitations", 2)
add_body(doc,
    "This study uses GPT-2 (124M), a small model; results may not generalise to larger "
    "LLMs. The reward model uses automatic IMDB labels rather than human pairwise "
    "preferences. Training is limited to 200 PPO steps on CPU hardware. DPO preference "
    "pairs are constructed automatically, potentially underestimating DPO's capability "
    "with richer human-annotated data.", indent=True)

# ═══════════════════════════════════════════════════════════════
# VI. CONCLUSION
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "VI. Conclusion", 1)
add_body(doc,
    f"This work presents a complete, reproducible RLHF pipeline on GPT-2/IMDB and draws "
    f"three actionable conclusions. First, PPO with β=0.20 reliably improves sentiment "
    f"reward by +{agg['reward_delta']:.2f} points (+{agg['alignment_tax_pct']:.1f}% ppl) "
    f"across all tested seeds. Second, PPO outperforms DPO in reward efficiency at this "
    f"scale (2.5× lower perplexity cost per reward gain). Third, β=0.20 is the "
    f"Pareto-optimal KL-penalty coefficient for GPT-2-scale RLHF, providing the best "
    f"stability-reward trade-off with a {opt['safety_margin_kl']:.1f}-unit KL safety margin. "
    f"We recommend β=0.20 as the default for future GPT-2-scale RLHF experiments, "
    f"with an early-stopping rule: halt if mean KL > {opt['collapse_onset_kl']:.1f}.", indent=True)

# ═══════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "References", 1)

refs = [
    "[1] P. F. Christiano, J. Leike, T. Brown, M. Martic, S. Legg, and D. Amodei, \"Deep reinforcement learning from human preferences,\" in Advances in Neural Information Processing Systems, vol. 30, 2017.",
    "[2] L. Ouyang et al., \"Training language models to follow instructions with human feedback,\" in Advances in Neural Information Processing Systems, vol. 35, pp. 27730–27744, 2022.",
    "[3] L. Gao et al., \"Scaling laws for reward model overoptimization,\" in Proc. Int. Conf. Machine Learning (ICML), 2023.",
    "[4] R. Rafailov, A. Sharma, E. Mitchell, C. D. Manning, S. Ermon, and C. Finn, \"Direct preference optimization: Your language model is secretly a reward model,\" in Advances in Neural Information Processing Systems, vol. 36, 2023.",
    "[5] Y. Bai et al., \"Training a helpful and harmless assistant with reinforcement learning from human feedback,\" arXiv preprint arXiv:2204.05862, 2022.",
    "[6] A. Radford, J. Wu, R. Child, D. Luan, D. Amodei, and I. Sutskever, \"Language models are unsupervised multitask learners,\" OpenAI Blog, vol. 1, no. 8, p. 9, 2019.",
    "[7] A. L. Maas et al., \"Learning word vectors for sentiment analysis,\" in Proc. 49th Annual Meeting of the Association for Computational Linguistics, pp. 142–150, 2011.",
    "[8] L. von Werra et al., \"TRL: Transformer reinforcement learning,\" GitHub, 2022. [Online]. Available: https://github.com/huggingface/trl",
    "[9] J. Schulman, F. Wolski, P. Dhariwal, A. Radford, and O. Klimov, \"Proximal policy optimization algorithms,\" arXiv preprint arXiv:1707.06347, 2017.",
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)

# ── Save ──────────────────────────────────────────────────────
out_path = os.path.join(OUT_DIR, "RLHF_IEEE_Report.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
